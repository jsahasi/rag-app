"""Document loading and text chunking utilities."""

import os
from pathlib import Path
from typing import Generator
from dataclasses import dataclass

from config import Config


@dataclass
class Document:
    """Represents a document chunk."""
    content: str
    metadata: dict

    @property
    def id(self) -> str:
        """Generate unique ID for the document chunk."""
        return f"{self.metadata['source']}_{self.metadata['chunk_index']}"


class DocumentLoader:
    """Loads documents from various file formats."""

    def __init__(self, folder_path: str):
        self.folder_path = Path(folder_path).resolve()
        self.supported_extensions = Config.get_supported_extensions()

    def load_all(self) -> list[Document]:
        """Load all supported documents from the folder."""
        documents = []
        for file_path in self._get_files():
            try:
                docs = self._load_file(file_path)
                documents.extend(docs)
            except Exception as e:
                print(f"Warning: Could not load {file_path}: {e}")
        return documents

    def _get_files(self) -> Generator[Path, None, None]:
        """Recursively get all supported files."""
        for root, _, files in os.walk(self.folder_path):
            root_path = Path(root)
            # Skip index folder
            if Config.INDEX_FOLDER in root_path.parts:
                continue
            for file in files:
                file_path = root_path / file
                if file_path.suffix.lower() in self.supported_extensions:
                    yield file_path

    def _load_file(self, file_path: Path) -> list[Document]:
        """Load a single file and return document chunks."""
        suffix = file_path.suffix.lower()

        if suffix in Config.TEXT_EXTENSIONS or suffix in Config.CODE_EXTENSIONS:
            content = self._load_text_file(file_path)
        elif suffix in Config.PDF_EXTENSIONS:
            content = self._load_pdf_file(file_path)
        elif suffix in Config.DOCX_EXTENSIONS:
            content = self._load_docx_file(file_path)
        else:
            return []

        if not content.strip():
            return []

        return self._chunk_text(content, file_path)

    def _load_text_file(self, file_path: Path) -> str:
        """Load a text or code file."""
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        for encoding in encodings:
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return ""

    def _load_pdf_file(self, file_path: Path) -> str:
        """Load a PDF file."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            print("Warning: pypdf not installed. Cannot load PDF files.")
            return ""
        except Exception as e:
            print(f"Warning: Could not read PDF {file_path}: {e}")
            return ""

    def _load_docx_file(self, file_path: Path) -> str:
        """Load a DOCX file."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            return "\n\n".join(text_parts)
        except ImportError:
            print("Warning: python-docx not installed. Cannot load DOCX files.")
            return ""
        except Exception as e:
            print(f"Warning: Could not read DOCX {file_path}: {e}")
            return ""

    def _chunk_text(self, text: str, file_path: Path) -> list[Document]:
        """Split text into overlapping chunks."""
        chunks = []
        chunk_size = Config.CHUNK_SIZE
        chunk_overlap = Config.CHUNK_OVERLAP

        # Get relative path from folder
        try:
            relative_path = file_path.relative_to(self.folder_path)
        except ValueError:
            relative_path = file_path.name

        start = 0
        chunk_index = 0

        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]

            if chunk_text.strip():
                chunks.append(Document(
                    content=chunk_text,
                    metadata={
                        "source": str(relative_path),
                        "chunk_index": chunk_index,
                        "file_type": file_path.suffix.lower(),
                    }
                ))

            start = end - chunk_overlap
            chunk_index += 1

        return chunks


def get_instructions(folder_path: str) -> str | None:
    """Load instructions.txt from the folder if it exists."""
    instructions_path = Path(folder_path) / Config.INSTRUCTIONS_FILE
    if instructions_path.exists():
        try:
            return instructions_path.read_text(encoding="utf-8")
        except Exception:
            return None
    return None
